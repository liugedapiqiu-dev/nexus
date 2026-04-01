#!/usr/bin/env python3
"""
批量导入文件到 Information Memory

支持：
- .xlsx
- .csv
- .docx
- .pdf
- 其他文本文件（兜底）
- .ai / .xls 等难解析格式先记录元信息，不强行灌脏文本

输出：
- ~/.vectorbrain/memory/information_memory.db
"""

import json
import math
import hashlib
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Iterable

import openpyxl
import docx
import csv
import sys

TOOLS_VENV = Path.home() / '.vectorbrain' / '.venv-tools'
for candidate in [
    TOOLS_VENV / 'lib' / 'python3.14' / 'site-packages',
    TOOLS_VENV / 'lib' / 'python3.13' / 'site-packages',
    TOOLS_VENV / 'lib' / 'python3.12' / 'site-packages',
    TOOLS_VENV / 'lib' / 'python3.11' / 'site-packages',
]:
    if candidate.exists() and str(candidate) not in sys.path:
        sys.path.insert(0, str(candidate))

try:
    import pypdf
except Exception:
    pypdf = None

VECTORBRAIN_ROOT = Path.home() / '.vectorbrain'
INFO_DB = VECTORBRAIN_ROOT / 'memory' / 'information_memory.db'


def ensure_db():
    conn = sqlite3.connect(INFO_DB)
    conn.execute(
        '''
        CREATE TABLE IF NOT EXISTS knowledge (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            category TEXT NOT NULL,
            key TEXT NOT NULL UNIQUE,
            value TEXT NOT NULL,
            source_worker TEXT,
            confidence REAL DEFAULT 1.0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            embedding_vector TEXT
        )
        '''
    )
    conn.execute('CREATE INDEX IF NOT EXISTS idx_info_category ON knowledge(category)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_info_key ON knowledge(key)')
    conn.execute('CREATE INDEX IF NOT EXISTS idx_info_updated_at ON knowledge(updated_at)')
    conn.commit()
    conn.close()


def get_ollama_embedding(text: str):
    import subprocess
    result = subprocess.run(['ollama', 'run', 'bge-m3', text], capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(result.stderr[:300])
    return json.loads(result.stdout.strip())


def stable_key(prefix: str, text: str) -> str:
    h = hashlib.md5(text.encode('utf-8', errors='ignore')).hexdigest()[:12]
    return f'{prefix}:{h}'


def chunk_text(text: str, max_chars: int = 2800) -> List[str]:
    text = text.strip()
    if not text:
        return []
    paras = [p.strip() for p in text.split('\n') if p.strip()]
    chunks, current = [], []
    current_len = 0
    for para in paras:
        if current_len + len(para) + 1 > max_chars and current:
            chunks.append('\n'.join(current))
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += len(para) + 1
    if current:
        chunks.append('\n'.join(current))
    return chunks


def normalize_cell(v):
    if v is None:
        return ''
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
    return str(v).strip()


def parse_xlsx(path: Path) -> List[Dict]:
    wb = openpyxl.load_workbook(path, data_only=True)
    records = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        lines = []
        for row in rows:
            vals = [normalize_cell(v) for v in row]
            if any(vals):
                lines.append(' | '.join(vals))
        sheet_text = '\n'.join(lines)
        if not sheet_text.strip():
            continue
        for idx, chunk in enumerate(chunk_text(sheet_text, max_chars=3200), 1):
            records.append({
                'category': 'information_memory.xlsx',
                'title': f'{path.name} / 工作表: {ws.title} / 分块 {idx}',
                'text': chunk,
                'meta': {
                    'file_path': str(path),
                    'file_name': path.name,
                    'sheet_name': ws.title,
                    'chunk_index': idx,
                    'file_type': 'xlsx',
                }
            })
    return records


def parse_docx(path: Path) -> List[Dict]:
    d = docx.Document(str(path))
    parts = []
    for p in d.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)
    for table in d.tables:
        for row in table.rows:
            vals = [cell.text.strip() for cell in row.cells]
            if any(vals):
                parts.append(' | '.join(vals))
    full_text = '\n'.join(parts)
    records = []
    for idx, chunk in enumerate(chunk_text(full_text, max_chars=3200), 1):
        records.append({
            'category': 'information_memory.docx',
            'title': f'{path.name} / 分块 {idx}',
            'text': chunk,
            'meta': {
                'file_path': str(path),
                'file_name': path.name,
                'chunk_index': idx,
                'file_type': 'docx',
            }
        })
    return records


def parse_csv(path: Path) -> List[Dict]:
    lines = []
    with open(path, 'r', encoding='utf-8', errors='ignore', newline='') as f:
        reader = csv.reader(f)
        for row in reader:
            vals = [normalize_cell(v) for v in row]
            if any(vals):
                lines.append(' | '.join(vals))
    text = '\n'.join(lines)
    return [{
        'category': 'information_memory.csv',
        'title': f'{path.name} / 分块 {idx}',
        'text': chunk,
        'meta': {
            'file_path': str(path),
            'file_name': path.name,
            'chunk_index': idx,
            'file_type': 'csv',
        }
    } for idx, chunk in enumerate(chunk_text(text, max_chars=3200), 1)]


def parse_pdf(path: Path) -> List[Dict]:
    if pypdf is None:
        return [{
            'category': 'information_memory.file_meta',
            'title': f'{path.name} / PDF待解析',
            'text': f'PDF 文件待解析（当前环境缺少 pypdf）：{path.name}',
            'meta': {
                'file_path': str(path),
                'file_name': path.name,
                'file_type': 'pdf',
                'status': 'pending_pdf_parser'
            }
        }]
    reader = pypdf.PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        txt = (page.extract_text() or '').strip()
        if txt:
            parts.append(f'[PAGE {i}]\n{txt}')
    full_text = '\n'.join(parts)
    if not full_text.strip():
        return [{
            'category': 'information_memory.file_meta',
            'title': f'{path.name} / PDF无可提取文本',
            'text': f'PDF 文件存在，但当前没有提取到可读文本：{path.name}',
            'meta': {
                'file_path': str(path),
                'file_name': path.name,
                'file_type': 'pdf',
                'status': 'no_extractable_text'
            }
        }]
    return [{
        'category': 'information_memory.pdf',
        'title': f'{path.name} / 分块 {idx}',
        'text': chunk,
        'meta': {
            'file_path': str(path),
            'file_name': path.name,
            'chunk_index': idx,
            'file_type': 'pdf',
        }
    } for idx, chunk in enumerate(chunk_text(full_text, max_chars=3200), 1)]


def parse_text_file(path: Path) -> List[Dict]:
    text = path.read_text(encoding='utf-8', errors='ignore')
    return [{
        'category': 'information_memory.file',
        'title': path.name,
        'text': chunk,
        'meta': {
            'file_path': str(path),
            'file_name': path.name,
            'chunk_index': idx,
            'file_type': path.suffix.lower() or 'file',
        }
    } for idx, chunk in enumerate(chunk_text(text, max_chars=3200), 1)]


def parse_file_meta(path: Path, reason: str) -> List[Dict]:
    return [{
        'category': 'information_memory.file_meta',
        'title': f'{path.name} / 文件元信息',
        'text': f'文件已发现，暂未深度解析。原因：{reason}',
        'meta': {
            'file_path': str(path),
            'file_name': path.name,
            'file_type': path.suffix.lower() or 'unknown',
            'status': 'metadata_only',
            'reason': reason,
            'size_bytes': path.stat().st_size,
        }
    }]


def iter_records(root: Path) -> Iterable[Dict]:
    for path in sorted(root.rglob('*')):
        if not path.is_file():
            continue
        if path.name.startswith('.'):
            continue
        suffix = path.suffix.lower()
        try:
            if suffix == '.xlsx':
                yield from parse_xlsx(path)
            elif suffix == '.docx':
                yield from parse_docx(path)
            elif suffix == '.csv':
                yield from parse_csv(path)
            elif suffix == '.pdf':
                yield from parse_pdf(path)
            elif suffix in {'.md', '.txt', '.json', '.jsonl'}:
                yield from parse_text_file(path)
            elif suffix in {'.ai', '.xls'}:
                yield from parse_file_meta(path, reason='当前阶段仅记录元信息，避免脏文本入库')
            else:
                yield from parse_file_meta(path, reason='暂未定义解析器，先记录元信息')
        except Exception as e:
            yield {
                'category': 'information_memory.import_error',
                'title': f'{path.name} / 导入失败',
                'text': f'导入失败: {e}',
                'meta': {
                    'file_path': str(path),
                    'file_name': path.name,
                    'file_type': suffix or 'unknown',
                    'error': str(e),
                }
            }


def upsert_record(conn, key: str, category: str, value: str, embedding_vector: str, source_worker: str = 'information_memory_importer', confidence: float = 1.0):
    now = datetime.now().isoformat()
    conn.execute(
        '''
        INSERT INTO knowledge (category, key, value, source_worker, confidence, created_at, updated_at, embedding_vector)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(key) DO UPDATE SET
            category=excluded.category,
            value=excluded.value,
            source_worker=excluded.source_worker,
            confidence=excluded.confidence,
            updated_at=excluded.updated_at,
            embedding_vector=excluded.embedding_vector
        ''',
        (category, key, value, source_worker, confidence, now, now, embedding_vector)
    )


def import_directory(root_path: str):
    ensure_db()
    root = Path(root_path).expanduser().resolve()
    conn = sqlite3.connect(INFO_DB)
    imported = 0
    files = set()
    categories = {}

    for rec in iter_records(root):
        payload = json.dumps({
            'title': rec['title'],
            'content': rec['text'],
            'meta': rec['meta'],
        }, ensure_ascii=False)
        key = stable_key(str(root.name) + '|' + rec['meta'].get('file_path', ''), rec['title'] + '\n' + rec['text'])
        try:
            emb = get_ollama_embedding((rec['title'] + '\n' + rec['text'])[:4000])
            emb_json = json.dumps(emb, ensure_ascii=False)
        except Exception as e:
            emb_json = json.dumps({'embedding_error': str(e)}, ensure_ascii=False)
        upsert_record(conn, key=key, category=rec['category'], value=payload, embedding_vector=emb_json)
        imported += 1
        files.add(rec['meta'].get('file_path', ''))
        categories[rec['category']] = categories.get(rec['category'], 0) + 1
        if imported % 10 == 0:
            conn.commit()
            print(f'IMPORTED {imported}')

    conn.commit()
    conn.close()
    print(json.dumps({
        'root': str(root),
        'file_count': len(files),
        'record_count': imported,
        'categories': categories,
        'db': str(INFO_DB),
    }, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print('usage: import_information_files.py <dir>')
        raise SystemExit(1)
    import_directory(sys.argv[1])
